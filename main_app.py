import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import io
import os
import base64

# --- 1. CONFIGURATION ---
# We keep the folders organized
SAVE_DIR = "past_papers"
MS_DIR = "pyp_anschm"

for folder in [SAVE_DIR, MS_DIR]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# --- 1. CONFIGURATION (Continued) ---
variants_map = {
    "1": ["11", "12", "13"],
    "2": ["02", "21", "22", "23"],
    "3": ["31", "32", "33"],
    "4": ["04", "41", "42", "43"]
}

# --- 2. TOOLS ---
def extract_questions(path, keyword):
    """Extracts text from PDF based on a keyword search."""
    if not os.path.exists(path): return None
    try:
        doc = fitz.open(path)
        output = ""
        for page in doc:
            text = page.get_text()
            if not keyword or keyword.lower() in text.lower():
                output += f"\n--- {os.path.basename(path)} (P.{page.number + 1}) ---\n{text}"
        doc.close()
        return output if output.strip() else None
    except Exception as e:
        return f"Error reading PDF: {e}"

def display_pdf(file_path):
    """Renders a PDF file directly in the Streamlit app."""
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

# --- 3. THE INTERFACE ---
st.set_page_config(page_title="9626 IT Lecturer Portal", layout="wide")
st.title("📂 IT 9626 Questions Retrieval Suite")

# SIDEBAR: ADMIN UPLOAD
with st.sidebar:
    st.header("📤 New PYP Upload")
    # Change from selectbox to number_input for an empty/typeable entry box
    u_y = st.number_input("Year", min_value=2019, max_value=2099, value=2026, step=1)
    u_s = st.selectbox("Session", ["MARCH", "JUNE", "NOVEMBER"])
    u_p = st.selectbox("Select Paper", ["1", "2", "3", "4"])

    variants_map = {"1": ["11", "12", "13"], "2": ["02", "21", "22", "23"], "3": ["31", "32", "33"],
                    "4": ["04", "41", "42", "43"]}
    #u_v = st.selectbox("Select Variant", variants_map[u_p])
    u_v = st.selectbox("Select Variant", variants_map[u_p])

    u_file = st.file_uploader("Upload PDF", type="pdf")
    u_type = st.radio("File Type", ["Question Paper (QP)", "Marking Scheme (MS)"])

#*********************************************************************
    if st.button("Save to Database"):
        if u_file:
            # 1. Map the session name to the letter (MARCH -> m)
            s_map = {"MARCH": "m", "JUNE": "s", "NOVEMBER": "w"}
            session_letter = s_map[u_s]
            # 2. Get the last two digits of the year (2023 -> 23)
            short_year = str(u_y)[-2:]
            # 3. Decide the folder and prefix based on File Type
            if u_type == "Question Paper (QP)":
                target_dir = SAVE_DIR
                prefix = "qp"
            else:
                target_dir = MS_DIR
                prefix = "ms"
            # 4. CONSTRUCT THE FILENAME CORRECTLY
            # This ensures it uses exactly what you selected in the boxes!
            new_filename = f"9626_{session_letter}{short_year}_{prefix}_{u_v}.pdf"
            save_path = os.path.join(target_dir, new_filename)

            # 5. Save the file
            with open(save_path, "wb") as f:
                f.write(u_file.getbuffer())

            # 6. Success Message (The Fix for the 'm26' vs 'm23' issue)
            st.success(f"✅ Successfully saved as: {new_filename}")

# MAIN TABS
t1, t2 = st.tabs(["🔍 One PYP Extractor", "📚 Batch PYP Extractor (Refresh)"])

with t1:
    st.subheader("🔍 Find Exam Resources")
    col1, col2, col3 = st.columns(3)

    with col1:
        s_topic = st.text_input("Topic Keyword", placeholder="e.g. network", key="st_topic")
    with col2:
        s_year = st.selectbox("Year", range(2019, 2025), index=0, key="st_year")
    with col3:
        s_p = st.selectbox("Paper", ["1", "2", "3", "4"], key="st_paper")
        s_v = st.selectbox("Variant", variants_map[s_p], key="st_variant")

    if st.button("Search & Extract", key="st_btn"):
        found = []
        for s in ["m", "s", "w"]:
            fn = f"9626_{s}{str(s_year)[-2:]}_qp_{s_v}.pdf"
            file_path = os.path.join(SAVE_DIR, fn)
            if os.path.exists(file_path):
                res = extract_questions(file_path, s_topic)
                if res: found.append(res)

        if found:
            st.session_state['current_data'] = found
            for item in found: st.info(item[:500] + "...")
        else:
            st.error(f"No questions found for Variant {s_v} in {s_year}.")

#for paper number aligned with variant number
with t2:
    st.subheader("📚 4-Years Topical Batch Extractor")
    st.write("Extract questions across a 4-year range into one document.")

    col_a, col_b = st.columns(2)
    with col_a:
        # Starting year input for the 4-year range
        start_yr = st.number_input("Starting Year", min_value=2019, max_value=2030, value=2019, key="batch_start")
        end_yr = start_yr + 3
        st.info(f"Scanning from {start_yr} to {end_yr}")

    with col_b:
        # DYNAMIC SELECTION: This is the change you requested
        b_paper = st.selectbox("Paper Number", ["1", "2", "3", "4"], key="batch_p")
        # Automatically filters based on the paper number chosen above
        b_variant = st.selectbox("Variant", variants_map[b_paper], key="batch_v")

    b_topic = st.text_input("Topic Keyword (e.g. 'Software')", key="batch_topic")

    if st.button("🚀 Run Batch Extraction", key="batch_run"):
        all_text = ""
        # The engine loops through the defined 4-year window
        for year in range(start_yr, end_yr + 1):
            for session in ["m", "s", "w"]:
                fn = f"9626_{session}{str(year)[-2:]}_qp_{b_variant}.pdf"
                path = os.path.join(SAVE_DIR, fn)

                if os.path.exists(path):
                    res = extract_questions(path, b_topic)
                    if res:
                        all_text += f"\n\n{'=' * 40}\nYEAR: {year} | SESSION: {session.upper()}\n{'=' * 40}\n"
                        all_text += res

        if all_text:
            st.success(f"Successfully compiled questions for {start_yr}-{end_yr}")
            st.text_area("Content Preview", all_text[:1000] + "...", height=250)

            # THE SAVE BUTTON
            doc = Document()
            doc.add_heading(f'Topical Question Bank: {b_topic}', 0)
            doc.add_paragraph(f"Range: {start_yr} to {end_yr} | Paper: {b_paper} | Variant: {b_variant}")
            doc.add_paragraph(all_text)

            bio = io.BytesIO()
            doc.save(bio)
            st.download_button(
                label="💾 SAVE AS WORD DOCUMENT",
                data=bio.getvalue(),
                file_name=f"Batch_{b_topic}_{start_yr}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning(f"No matching questions found in the {start_yr}-{end_yr} window.")

# ACTION FOOTER
if 'current_data' in st.session_state:
    st.divider()
    c_view, c_doc = st.columns(2)

    if c_view.button("🎯 View Official Marking Scheme"):
        ms_found = False
        for item in st.session_state['current_data']:
            if "---" in item:
                #qp_name = item.split("---")[1].strip()
                # This strips away the " (P.x)" part by cutting the text at ".pdf"
                qp_name = item.split("---")[1].split(".pdf")[0].strip() + ".pdf"
                ms_name = qp_name.replace("_qp_", "_ms_")
                ms_path = os.path.join(MS_DIR, ms_name)

                if os.path.exists(ms_path):
                    st.subheader(f"📄 Displaying: {ms_name}")
                    display_pdf(ms_path)
                    ms_found = True
                else:
                    st.warning(f"Official Scheme {ms_name} is not in the '{MS_DIR}' folder yet.")

        if not ms_found:
            st.error("No marking scheme files found matching your search.")

    if c_doc.button("💾 Download as Word Document"):
        doc = Document()
        doc.add_heading('9626 IT Extracted Questions', 0)
        for text in st.session_state['current_data']:
            doc.add_paragraph(text)
        bio = io.BytesIO()
        doc.save(bio)

        st.download_button("Click to Download", bio.getvalue(), "IT_Questions.docx")
